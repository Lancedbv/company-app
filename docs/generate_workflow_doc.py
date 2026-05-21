from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(30, 30, 50)

for s in ['Heading 1', 'Heading 2', 'Heading 3']:
    hs = doc.styles[s]
    hs.font.color.rgb = RGBColor(96, 77, 255)
    hs.font.name = 'Calibri'

# ── Title ──
title = doc.add_heading('Lanced — Opportunity Types, Formats & Workflows', level=0)
title.runs[0].font.color.rgb = RGBColor(30, 30, 50)
doc.add_paragraph('Complete reference for all opportunity types, access formats, workflow stages, and feature modules.\nLast updated: May 2026')

# ═══════════════════════════════════════
# SECTION 1: Opportunity Types
# ═══════════════════════════════════════
doc.add_heading('1. Opportunity Types', level=1)
doc.add_paragraph('Lanced supports 7 distinct opportunity types. Each type determines available features, workflow options, and candidate journey.')

types_data = [
    ('Casting', '#604DFF', 'Cast roles for a campaign, production or shoot', 'Short-term project hiring. Artists apply or are scouted, reviewed, and cast for specific roles in productions, campaigns, or shoots.'),
    ('Audition', '#1A56DB', 'Run season auditions for company members or roles', 'Multi-round selection process. Supports group/individual formats, registration, live scoring, and progression through callback rounds to final offers.'),
    ('Job Call', '#0EA5A8', 'Hire for staff or freelance roles within your company', 'Traditional hiring flow. Includes application review, interviews (auto-enabled), and contract offers. Best for permanent or freelance staff positions.'),
    ('Open Call', '#F5A623', 'Public call to discover new artists and talent', 'Broad talent discovery. Maximizes reach — anyone can apply. Supports interviews and large-volume screening. Ideal for scouting new talent.'),
    ('Residency', '#FF6B81', 'Invite artists for a creative or research residency', 'Time-bound artist programs. Includes interview scheduling and contract management for hosted creative periods.'),
    ('Competition', '#1DB954', 'Run a contest, prize or competitive selection', 'Scored evaluation. Enables custom scoring criteria with weighted percentages, leaderboards, and ranked results.'),
    ('Education', '#8B5CF6', 'Run a school, training programme or workshop', 'Educational programs. Supports contracts for enrollment terms and structured group sessions.'),
]

table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Type'
hdr[1].text = 'Description'
hdr[2].text = 'Key Features'
hdr[3].text = 'Special Modules'

special_modules = [
    'Scoring (optional)',
    'Multi-round, Group/Individual formats, Registration, Contracts',
    'Interviews (auto-enabled), Contracts',
    'Interviews (eligible), High-volume screening',
    'Interviews (eligible), Contracts',
    'Scoring (auto-enabled), Leaderboard',
    'Contracts',
]

for i, (name, color, desc, detail) in enumerate(types_data):
    row = table.add_row().cells
    row[0].text = name
    row[1].text = desc
    row[2].text = detail
    row[3].text = special_modules[i]

# ═══════════════════════════════════════
# SECTION 2: Access Formats
# ═══════════════════════════════════════
doc.add_heading('2. Access Formats', level=1)
doc.add_paragraph('Each opportunity can be configured with different access levels, controlling who can view and apply.')

doc.add_heading('Standard Access Types (all non-audition types)', level=2)
table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'
hdr2 = table2.rows[0].cells
hdr2[0].text = 'Format'
hdr2[1].text = 'Label'
hdr2[2].text = 'Description'

access_std = [
    ('open', 'Open Call', 'Public — anyone can view the listing and submit an application'),
    ('closed_link', 'By Invite Link', 'Private — share a unique link with selected candidates to apply'),
    ('closed_internal', 'Internal Only', 'Private — select candidates directly from your artist database, no external applications'),
]
for key, label, desc in access_std:
    row = table2.add_row().cells
    row[0].text = key
    row[1].text = label
    row[2].text = desc

doc.add_heading('Audition Access Types', level=2)
table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Light Grid Accent 1'
hdr3 = table3.rows[0].cells
hdr3[0].text = 'Format'
hdr3[1].text = 'Label'
hdr3[2].text = 'Description'

access_aud = [
    ('open', 'Open Audition', 'Anyone can apply — applicants auto-join the audition pipeline'),
    ('closed', 'Closed Audition', 'Invite-only — company reviews applications before granting access'),
]
for key, label, desc in access_aud:
    row = table3.add_row().cells
    row[0].text = key
    row[1].text = label
    row[2].text = desc

doc.add_heading('Audition Formats', level=2)
table3b = doc.add_table(rows=1, cols=3)
table3b.style = 'Light Grid Accent 1'
hdr3b = table3b.rows[0].cells
hdr3b[0].text = 'Format'
hdr3b[1].text = 'Title'
hdr3b[2].text = 'Description'

aud_formats = [
    ('private', 'Private Auditions', 'Individual time-slot booking — artists pick their own audition slot (Calendly-style)'),
    ('regular', 'Regular Audition', 'Everyone invited to same event — no groups, single session'),
    ('groups', 'Group Audition', 'Artists divided into groups — by self-booking, company assignment, or on the day'),
    ('multi_location', 'Multi-Location', 'Auditions across cities/studios — each location runs independently'),
]
for key, title, desc in aud_formats:
    row = table3b.add_row().cells
    row[0].text = key
    row[1].text = title
    row[2].text = desc

# ═══════════════════════════════════════
# SECTION 3: Workflow Stages
# ═══════════════════════════════════════
doc.add_heading('3. Workflow & Pipeline Stages', level=1)

doc.add_heading('Candidate Lifecycle (Room-level)', level=2)
doc.add_paragraph('Every applicant moves through these statuses at the opportunity level:')

table4 = doc.add_table(rows=1, cols=3)
table4.style = 'Light Grid Accent 1'
hdr4 = table4.rows[0].cells
hdr4[0].text = 'Status'
hdr4[1].text = 'Label'
hdr4[2].text = 'Description'

statuses = [
    ('new', 'New', 'Just applied — unreviewed'),
    ('potential', 'Potential', 'Under consideration, not yet shortlisted'),
    ('shortlisted', 'Shortlisted', 'Marked as strong candidate (when Shortlist is enabled)'),
    ('waitlisted', 'Waitlisted', 'On waitlist — backup if selected candidates decline (when Waitlist is enabled)'),
    ('selected', 'Selected / Offered', 'Chosen for the role — offer extended'),
    ('not_selected', 'Not Selected', 'Declined / rejected from the process'),
    ('callback', 'Callback', 'Advanced to the next round (audition-specific)'),
]
for key, label, desc in statuses:
    row = table4.add_row().cells
    row[0].text = key
    row[1].text = label
    row[2].text = desc

doc.add_heading('Round Phases (Audition-specific)', level=2)
doc.add_paragraph('Each audition round progresses through 4 phases:')

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Light Grid Accent 1'
hdr5 = table5.rows[0].cells
hdr5[0].text = 'Phase'
hdr5[1].text = 'Label'
hdr5[2].text = 'What happens'

phases = [
    ('setup', 'Setup', 'Configure round: set groups, sessions, capacity, rules, dates, locations. Send invitations.'),
    ('registration', 'Registration', 'Day-of check-in. Mark artists as checked_in, late, or no_show. Edit numbers, reassign groups.'),
    ('audition', 'Audition (Live)', 'The round is live. Team votes, takes notes, captures photos/video. Real-time collaboration.'),
    ('decision', 'Decision', 'Bulk-decide outcomes: Callback (advance) or Not Selected. Transition to next round.'),
]
for key, label, desc in phases:
    row = table5.add_row().cells
    row[0].text = key
    row[1].text = label
    row[2].text = desc

doc.add_heading('Confirmation Flow (after selection/invitation)', level=2)
p = doc.add_paragraph()
p.add_run('When a candidate is invited to a round or selected for a role, they go through:\n').bold = False
confirm = ['pending → Awaiting response', 'confirmed → Accepted', 'declined → Declined the invitation', 'rescheduling → Requested a different time']
for c in confirm:
    doc.add_paragraph(c, style='List Bullet')

# ═══════════════════════════════════════
# SECTION 4: Feature Modules
# ═══════════════════════════════════════
doc.add_heading('4. Feature Modules (Toggle On/Off)', level=1)
doc.add_paragraph('These features can be enabled or disabled per opportunity to customize the workflow.')

doc.add_heading('Shortlist', level=2)
doc.add_paragraph('Purpose: Create a curated list of top candidates before making final selections.\n\nHow it works: After reviewing applications, mark promising candidates as "Shortlisted." This creates an intermediate tier between new applicants and selected candidates. The shortlist serves as a working list for the team to discuss and narrow down before committing to offers.\n\nWhen to use: When you receive many applications and need a way to separate "definitely consider" from "maybe" before making final decisions. Especially useful for large open calls and castings.')

doc.add_heading('Waitlist', level=2)
doc.add_paragraph('Purpose: Maintain a ranked backup list of candidates in case selected candidates decline.\n\nHow it works: When a shortlisted or considered candidate isn\'t selected but is still a strong fit, place them on the waitlist. If a selected candidate declines their offer, waitlisted candidates can be promoted to selected — preserving the pipeline without restarting the search.\n\nWhen to use: For competitive opportunities where you expect some selected candidates may decline. Common in auditions and residencies where artists may have conflicting commitments.')

doc.add_heading('Early Invite', level=2)
doc.add_paragraph('Purpose: Send advance invitations to priority candidates before opening general applications.\n\nHow it works: Before publishing the opportunity publicly (or before a round begins), send early invitations to specific artists from your database. These artists get first access to apply, book slots, or confirm participation. Their early invite status is tracked separately (pending → confirmed → declined).\n\nWhen to use: When you have returning talent or priority relationships. Common in season auditions where principal dancers or known artists get early access before the open call.')

doc.add_heading('Interview', level=2)
doc.add_paragraph('Purpose: Schedule structured 1:1 or panel conversations with candidates.\n\nHow it works: Enable interview scheduling with configurable duration, buffer time, and availability windows. Supports both specific dates and weekly recurring hours. Candidates can self-book available slots (Calendly-style) or be assigned times. Supports in-person and video (Google Meet) formats. Interview reminders can be sent 24h and 1h before.\n\nWhen to use: Auto-enabled for Job Calls. Also available for Open Calls and Residencies. Best for roles requiring conversation beyond portfolio review — discussing availability, terms, creative vision, etc.\n\nEligible types: Job Call (auto), Open Call, Residency')

doc.add_heading('Hire / Contracts', level=2)
doc.add_paragraph('Purpose: Formalize the selection with contract offers and track acceptance.\n\nHow it works: When a candidate is selected, attach a contract type (Full-time, Part-time, Freelance, Project-based, Internship, Apprenticeship, Seasonal). The offer is sent and the candidate confirms or declines. This is the terminal step of the pipeline — moving from "selected" to "hired."\n\nWhen to use: For any opportunity that results in a formal engagement. Available for Auditions, Job Calls, Residencies, and Education programs.\n\nContract types: Full Time, Part Time, Freelance, Contract, Internship, Apprenticeship, Project-Based, Seasonal')

doc.add_heading('Scoring', level=2)
doc.add_paragraph('Purpose: Evaluate candidates with structured, weighted scoring criteria.\n\nHow it works: Define custom criteria (e.g., Technique 30%, Creativity 30%, Artistry 20%, Stage Presence 20%) with configurable weights that must total 100%. Each team member scores candidates on a numeric scale. Scores are aggregated into a weighted total and displayed on a leaderboard for objective comparison.\n\nWhen to use: Auto-enabled for Competitions. Optional for all other types. Best when you need objective, comparable evaluations across many candidates.')

doc.add_heading('Voting', level=2)
doc.add_paragraph('Purpose: Quick team consensus on candidates during review.\n\nHow it works: Team members cast Yes / Maybe / No votes on each candidate. Votes are visible in real-time and inform (but don\'t auto-decide) candidate outcomes. Provides a lightweight way to gather team opinions without formal scoring.\n\nWhen to use: During live audition rounds or application review when the team needs to quickly signal preferences.')

doc.add_heading('Batches', level=2)
doc.add_paragraph('Purpose: Process candidates in grouped batches for efficient bulk review.\n\nHow it works: Divide candidates into manageable batches for systematic review. Process one batch at a time rather than reviewing the full list at once.\n\nWhen to use: For high-volume opportunities with 100+ applicants where reviewing all at once is impractical.')

# ═══════════════════════════════════════
# SECTION 5: Workflow Examples
# ═══════════════════════════════════════
doc.add_heading('5. Example Workflows', level=1)

doc.add_heading('A. Open Audition — Group Format', level=2)
doc.add_paragraph(
    '1. Create Opportunity → Type: Audition, Access: Open, Format: Group Audition\n'
    '2. Configure Round 1 → Preset: Multi-Group, set sessions with capacity/dates/locations\n'
    '3. Publish → Artists apply freely\n'
    '4. (Optional) Early Invite → Send priority invites to known talent\n'
    '5. Registration Phase → Day-of check-in, assign numbers\n'
    '6. Audition Phase → Live: team votes, notes, photos/video\n'
    '7. Decision Phase → Mark Callback or Not Selected\n'
    '8. Round 2 (Callback) → Smaller group, connect to previous round groups\n'
    '9. Finals → Single group, final decisions\n'
    '10. Select → Send offers with contracts\n'
    '11. Hire → Candidates confirm, contracts signed'
)

doc.add_heading('B. Closed Audition — Private (Individual)', level=2)
doc.add_paragraph(
    '1. Create Opportunity → Type: Audition, Access: Closed, Format: Private Auditions\n'
    '2. Early Invite → Send invitations to selected artists from database\n'
    '3. Artists self-book individual time slots (Calendly-style)\n'
    '4. Audition Phase → 1:1 sessions, team evaluates each artist\n'
    '5. Decision → Shortlist → Select → Contract offer'
)

doc.add_heading('C. Casting — Open Call', level=2)
doc.add_paragraph(
    '1. Create Opportunity → Type: Casting, Access: Open\n'
    '2. Publish → Artists apply with portfolio\n'
    '3. Review applications → Shortlist top candidates\n'
    '4. (Optional) Waitlist backup candidates\n'
    '5. Select → Notify selected candidates\n'
    '6. Candidates confirm or decline\n'
    '7. If decline → Promote from waitlist'
)

doc.add_heading('D. Job Call — Internal', level=2)
doc.add_paragraph(
    '1. Create Opportunity → Type: Job Call, Access: Internal Only\n'
    '2. Select candidates from artist database\n'
    '3. Interview (auto-enabled) → Schedule 1:1 interviews\n'
    '4. Candidates self-book slots or are assigned times\n'
    '5. Conduct interviews → Vote/score\n'
    '6. Select → Send contract offer (Full-time/Freelance/etc.)\n'
    '7. Hire → Confirmed'
)

doc.add_heading('E. Competition — Open', level=2)
doc.add_paragraph(
    '1. Create Opportunity → Type: Competition, Access: Open\n'
    '2. Define scoring criteria + weights (e.g., Technique 30%, Creativity 30%)\n'
    '3. Artists apply\n'
    '4. Round 1 → Groups perform, jury scores each candidate\n'
    '5. Leaderboard auto-ranks by weighted total\n'
    '6. Decision → Top scorers advance (Callback)\n'
    '7. Finals → Final scoring round\n'
    '8. Winner announced from leaderboard'
)

doc.add_heading('F. Residency — By Invite Link', level=2)
doc.add_paragraph(
    '1. Create Opportunity → Type: Residency, Access: By Invite Link\n'
    '2. Share private link with curated artists\n'
    '3. Artists apply through link\n'
    '4. Review → Shortlist\n'
    '5. Interview eligible candidates\n'
    '6. Select → Contract offer (dates, terms)\n'
    '7. Hire → Confirmed'
)

# ═══════════════════════════════════════
# SECTION 6: Round Presets
# ═══════════════════════════════════════
doc.add_heading('6. Round Presets', level=1)

table6 = doc.add_table(rows=1, cols=4)
table6.style = 'Light Grid Accent 1'
hdr6 = table6.rows[0].cells
hdr6[0].text = 'Preset'
hdr6[1].text = 'Label'
hdr6[2].text = 'Best For'
hdr6[3].text = 'Description'

presets = [
    ('multi_group', 'Multi-Group', '80–200+ artists', 'Groups run parallel or back-to-back with capacity limits. Artists self-book or are assigned.'),
    ('single_group', 'Single Group', 'Smaller callbacks', 'Everyone performs together in one session. No group division needed.'),
    ('interview', 'Interview', '1:1 or panel', 'Structured individual sessions with time-slotted scheduling.'),
    ('solo_audition', 'Solo Audition', 'Musicians/soloists', 'Single artist per scheduled slot. Individual performance evaluation.'),
    ('finals', 'Finals', 'Final decisions', 'Terminal round that triggers contract/offer flow. Last round in the process.'),
]
for key, label, best, desc in presets:
    row = table6.add_row().cells
    row[0].text = key
    row[1].text = label
    row[2].text = best
    row[3].text = desc

# ═══════════════════════════════════════
# SECTION 7: Session Rules
# ═══════════════════════════════════════
doc.add_heading('7. Session / Group Rules', level=1)
doc.add_paragraph('Each group or session within a round can have filter rules to auto-sort candidates:')

table7 = doc.add_table(rows=1, cols=3)
table7.style = 'Light Grid Accent 1'
hdr7 = table7.rows[0].cells
hdr7[0].text = 'Criterion'
hdr7[1].text = 'Type'
hdr7[2].text = 'Options'

rules = [
    ('Gender', 'Multi-select', 'Female, Male, Non-binary'),
    ('Role', 'Select', 'Ensemble, Soloist, First Soloist, Principal, Apprentice, Any'),
    ('Discipline', 'Select', 'Classical, Contemporary, Modern, Hip-Hop, Jazz, Tap, Any'),
    ('Age', 'Select', '16-18, 18-21, 21-25, 25-30, 30-35, 35-40, 40+, Any'),
    ('Contract', 'Select', 'Full-time, Part-time, Freelance, Project-based, Internship, Any'),
    ('Experience', 'Select', 'Student, 1-3 years, 3-5 years, 5-10 years, 10+ years, Any'),
    ('Height', 'Range', 'Numeric min–max'),
    ('Nationality', 'Text', 'Free text input'),
]
for crit, typ, opts in rules:
    row = table7.add_row().cells
    row[0].text = crit
    row[1].text = typ
    row[2].text = opts

# Save
doc.save('/Users/woutervertogen/Claude/lanced-company/docs/Lanced_Workflows_Overview.docx')
print('Word document saved successfully.')
